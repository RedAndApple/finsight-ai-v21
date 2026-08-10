from __future__ import annotations

import html
import io
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


BRAND_DARK = "0E4B4B"
BRAND_TEAL = "246D6D"
BRAND_LIGHT = "EAF3F3"
BRAND_RED = "C84242"
BRAND_GOLD = "B77A16"
TEXT_DARK = "213333"


def _years(result: dict[str, Any]) -> list[str]:
    years = {
        str(year)
        for row in result.get("financial_metrics", {}).values()
        for year in row.get("values", {})
        if str(year).isdigit()
    }
    return sorted(years, key=int)


def _fmt(value: Any) -> str:
    if value is None:
        return "—"
    try:
        number = float(value)
        if number.is_integer():
            return f"{int(number):,}".replace(",", " ")
        return f"{number:,.2f}".replace(",", " ")
    except (TypeError, ValueError):
        return str(value)


def _item_text(item: Any) -> str:
    if isinstance(item, dict):
        for key in ("action", "text", "observation", "finding", "title", "description"):
            value = item.get(key)
            if value:
                return str(value).strip()
        return "; ".join(f"{key}: {value}" for key, value in item.items() if value not in (None, "", [], {}))
    return str(item or "").strip()


def _items(values: Iterable[Any] | None, empty: str = "Нет подтвержденных выводов.") -> list[str]:
    output = [_item_text(item) for item in (values or [])]
    output = [item for item in output if item]
    return output or [empty]


def _analysis_sections() -> tuple[tuple[str, str], ...]:
    return (
        ("Сильные стороны", "strengths"),
        ("Зоны внимания", "weaknesses"),
        ("Существенные риски", "risks"),
        ("Рекомендованные действия", "management_actions"),
        ("Стратегические наблюдения", "strategic_observations"),
        ("ESG-наблюдения", "esg_observations"),
        ("Ограничения данных", "data_limitations"),
    )


def _style_excel_table(sheet, widths: dict[int, float] | None = None) -> None:
    thin = Side(style="thin", color="D3E1E1")
    sheet.sheet_view.showGridLines = False
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    sheet.row_dimensions[1].height = 32
    for cell in sheet[1]:
        cell.font = Font(bold=True, color="FFFFFF", size=10)
        cell.fill = PatternFill("solid", fgColor=BRAND_DARK)
        cell.alignment = Alignment(wrap_text=True, vertical="center")
    for row_number, row_cells in enumerate(sheet.iter_rows(), start=1):
        if row_number > 1 and row_number % 2 == 1:
            for cell in row_cells:
                cell.fill = PatternFill("solid", fgColor="F6FAFA")
        for cell in row_cells:
            cell.border = Border(bottom=thin)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    for index in range(1, sheet.max_column + 1):
        if widths and index in widths:
            width = widths[index]
        else:
            values = [str(sheet.cell(row, index).value or "") for row in range(1, min(sheet.max_row, 100) + 1)]
            width = min(50, max(12, max(map(len, values), default=12) + 2))
        sheet.column_dimensions[get_column_letter(index)].width = width
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 0
    sheet.sheet_properties.tabColor = BRAND_TEAL


def build_xlsx(result: dict[str, Any]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Резюме"
    metadata, analysis = result.get("metadata", {}), result.get("analysis", {})
    ws.sheet_view.showGridLines = False
    ws.merge_cells("A1:F2")
    ws["A1"] = "FinSight AI — профессиональный финансовый анализ"
    ws["A1"].font = Font(size=20, bold=True, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor=BRAND_DARK)
    ws["A1"].alignment = Alignment(vertical="center")
    ws.row_dimensions[1].height = 30
    ws.row_dimensions[2].height = 15
    summary_rows = [
        ("Организация", metadata.get("company")),
        ("Исходный файл", metadata.get("filename")),
        ("Стандарт", metadata.get("accounting_standard")),
        ("Отчетный год", metadata.get("reporting_year")),
        ("Периметр", metadata.get("reporting_scope")),
        ("Единица", metadata.get("unit") or metadata.get("currency")),
        ("Статус валидации", result.get("validation", {}).get("status")),
        ("Дата формирования", datetime.now().astimezone().strftime("%d.%m.%Y %H:%M")),
        ("Итоговое резюме", analysis.get("executive_summary")),
    ]
    for row, (label, value) in enumerate(summary_rows, start=4):
        ws.cell(row, 1, label).font = Font(bold=True, color=BRAND_DARK)
        ws.cell(row, 2, value or "—")
        ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
        ws.cell(row, 2).alignment = Alignment(wrap_text=True, vertical="top")
        if label == "Итоговое резюме":
            ws.row_dimensions[row].height = 75
    row = 15
    for title, key in _analysis_sections():
        ws.cell(row, 1, title).font = Font(size=12, bold=True, color="FFFFFF")
        ws.cell(row, 1).fill = PatternFill("solid", fgColor=BRAND_TEAL)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        row += 1
        for item in _items(analysis.get(key)):
            ws.cell(row, 1, f"• {item}")
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
            ws.cell(row, 1).alignment = Alignment(wrap_text=True, vertical="top")
            ws.row_dimensions[row].height = 36
            row += 1
        row += 1
    for column, width in {"A": 25, "B": 22, "C": 16, "D": 16, "E": 16, "F": 18}.items():
        ws.column_dimensions[column].width = width
    ws.sheet_properties.tabColor = BRAND_DARK
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0

    years = _years(result)
    fm = wb.create_sheet("Финансовая модель")
    fm.append(["Код", "Показатель", *years, "Единица", "Страницы", "Метод", "Доверие"])
    for item in result.get("financial_metrics", {}).values():
        fm.append([
            item.get("row_code") or "",
            item.get("name"),
            *[item.get("values", {}).get(year) for year in years],
            item.get("unit") or "",
            ", ".join(map(str, item.get("source_pages", []))),
            item.get("source_type") or "",
            float(item.get("confidence", 0) or 0),
        ])
    for row_cells in fm.iter_rows(min_row=2, min_col=3, max_col=2 + len(years)):
        for cell in row_cells:
            cell.number_format = '#,##0.00;[Red]-#,##0.00;—'
    for cell in fm[get_column_letter(fm.max_column)][1:]:
        cell.number_format = "0%"
    _style_excel_table(fm, {1: 11, 2: 31, **{3 + index: 17 for index in range(len(years))}})

    ratios = wb.create_sheet("Коэффициенты")
    ratios.append(["Коэффициент", "Значение", "Статус", "Формула", "Интерпретация"])
    for item in result.get("ratios", []):
        ratios.append([item.get("name"), item.get("display"), item.get("status"), item.get("formula"), item.get("explanation")])
    _style_excel_table(ratios, {1: 28, 2: 18, 3: 15, 4: 42, 5: 58})

    validation = wb.create_sheet("Валидация")
    validation.append(["Проверка", "Год", "Статус", "Левая часть", "Правая часть", "Отклонение", "Комментарий"])
    for item in result.get("validation", {}).get("checks", []):
        validation.append([
            item.get("name"), item.get("year"), item.get("status"), item.get("left"),
            item.get("right"), item.get("delta"), item.get("message") or item.get("issue"),
        ])
    _style_excel_table(validation, {1: 38, 2: 11, 3: 14, 4: 18, 5: 18, 6: 18, 7: 50})

    sources = wb.create_sheet("Источники")
    sources.append(["Показатель", "Страница / лист", "Строка / код", "Фрагмент", "Метод", "Доверие"])
    for item in result.get("financial_metrics", {}).values():
        provenance = item.get("provenance") or [{}]
        for source in provenance:
            sources.append([
                item.get("name"), source.get("page") or source.get("sheet"),
                source.get("row") or source.get("source_row") or source.get("row_code") or item.get("row_code"),
                source.get("text") or source.get("raw_text") or "",
                source.get("extraction_method") or item.get("source_type"),
                source.get("confidence", item.get("confidence")),
            ])
    for cell in sources["F"][1:]:
        cell.number_format = "0%"
    _style_excel_table(sources, {1: 30, 2: 17, 3: 26, 4: 60, 5: 30, 6: 12})

    notes = wb.create_sheet("Методология")
    notes.append(["Раздел", "Описание"])
    methodology = [
        ("Извлечение", "Текстовый слой, координатные таблицы, OCR и визуальное восстановление ключевых форм."),
        ("Нормализация", "Показатели приводятся к единой финансовой модели РСБУ/МСФО с сохранением страниц и строк."),
        ("Валидация", "Проверяются балансовые равенства, периоды, единицы и согласованность исходных строк."),
        ("Расчеты", "Коэффициенты и тренды рассчитываются программным кодом, а не языковой моделью."),
        ("AI-анализ", "Модель получает только нормализованные показатели, рассчитанные коэффициенты и проверенные раскрытия."),
        ("Ограничение", "Автоматический анализ не является аудиторским заключением или инвестиционной рекомендацией."),
    ]
    for entry in methodology:
        notes.append(entry)
    for limitation in result.get("limitations", []):
        notes.append(("Ограничение документа", _item_text(limitation)))
    _style_excel_table(notes, {1: 27, 2: 105})

    wb.calculation.fullCalcOnLoad = True
    wb.calculation.forceFullCalc = True
    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def _set_repeat_table_layout(table, widths_cm: list[float] | None = None) -> None:
    table.autofit = False
    _repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        _keep_row_together(row)
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell)
            if widths_cm and col_index < len(widths_cm):
                cell.width = Cm(widths_cm[col_index])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(8)
            if row_index == 0:
                _set_cell_shading(cell, BRAND_DARK)
                for run in cell.paragraphs[0].runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            elif row_index % 2 == 0:
                _set_cell_shading(cell, "F2F7F7")


def _docx_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    run.font.size = Pt(8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def _add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(4)


def _add_docx_section(document: Document, title: str, items: Iterable[Any] | None) -> None:
    document.add_heading(title, level=2)
    for item in _items(items):
        _add_bullet(document, item)


def build_docx(result: dict[str, Any]) -> bytes:
    metadata, analysis = result.get("metadata", {}), result.get("analysis", {})
    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(17)
    section.bottom_margin = Mm(17)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(TEXT_DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for style_name, size, color in (("Title", 25, BRAND_DARK), ("Heading 1", 18, BRAND_DARK), ("Heading 2", 13, BRAND_TEAL)):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
    if "FinSight Kicker" not in [style.name for style in styles]:
        kicker = styles.add_style("FinSight Kicker", WD_STYLE_TYPE.PARAGRAPH)
        kicker.font.name = "Aptos"
        kicker.font.size = Pt(9)
        kicker.font.bold = True
        kicker.font.color.rgb = RGBColor.from_string(BRAND_TEAL)
        kicker.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "FINSIGHT AI   |   ФИНАНСОВЫЙ УНИВЕРСИТЕТ"
    header.style = styles["FinSight Kicker"]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.add_run("FinSight AI · автоматический аналитический отчет   ")
    _docx_page_field(footer)

    document.add_paragraph("ПРОВЕРЯЕМЫЙ ФИНАНСОВЫЙ АНАЛИЗ", style="FinSight Kicker")
    document.add_heading(metadata.get("company") or "Финансовый отчет", level=0)
    subtitle = document.add_paragraph("РСБУ / МСФО · финансовая модель, коэффициенты, риски и рекомендации")
    subtitle.runs[0].font.size = Pt(12)
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(BRAND_TEAL)

    facts = document.add_table(rows=0, cols=2)
    facts.style = "Light Shading Accent 1"
    for label, value in (
        ("Исходный файл", metadata.get("filename")),
        ("Стандарт и периметр", " · ".join(filter(None, [str(metadata.get("accounting_standard") or ""), str(metadata.get("reporting_scope") or "")] ))),
        ("Отчетный год", metadata.get("reporting_year")),
        ("Статус валидации", result.get("validation", {}).get("status")),
        ("Сформировано", datetime.now().astimezone().strftime("%d.%m.%Y %H:%M")),
    ):
        cells = facts.add_row().cells
        cells[0].text = label
        cells[1].text = str(value or "—")
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(BRAND_DARK)
    _set_repeat_table_layout(facts, [4.5, 12.5])

    document.add_heading("Итоговое резюме", level=1)
    document.add_paragraph(analysis.get("executive_summary") or "Недостаточно подтвержденных данных для итогового резюме.")
    for title, key in _analysis_sections():
        _add_docx_section(document, title, analysis.get(key))

    document.add_page_break()
    document.add_heading("Финансовая модель", level=1)
    years = _years(result)
    headers = ["Код", "Показатель", *years, "Единица", "Источник"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = value
    for item in result.get("financial_metrics", {}).values():
        row = table.add_row().cells
        values = [
            item.get("row_code") or "", item.get("name"),
            *[_fmt(item.get("values", {}).get(year)) for year in years],
            item.get("unit") or "", ", ".join(map(str, item.get("source_pages", []))) or "—",
        ]
        for index, value in enumerate(values):
            row[index].text = str(value or "—")
            if 2 <= index < 2 + len(years):
                row[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    year_widths = [1.4, 5.4, *([2.1] * len(years)), 2.0, 2.0]
    _set_repeat_table_layout(table, year_widths)

    document.add_heading("Финансовые коэффициенты", level=1)
    ratio_headers = ["Коэффициент", "Значение", "Статус", "Формула", "Интерпретация"]
    ratio_table = document.add_table(rows=1, cols=len(ratio_headers))
    ratio_table.style = "Table Grid"
    for index, value in enumerate(ratio_headers):
        ratio_table.rows[0].cells[index].text = value
    for item in result.get("ratios", []):
        cells = ratio_table.add_row().cells
        values = [item.get("name"), item.get("display"), item.get("status"), item.get("formula"), item.get("explanation")]
        for index, value in enumerate(values):
            cells[index].text = str(value or "—")
    _set_repeat_table_layout(ratio_table, [3.2, 2.1, 1.8, 4.2, 5.7])

    document.add_heading("Валидация и ограничения", level=1)
    document.add_paragraph(f"Статус автоматической валидации: {result.get('validation', {}).get('status') or '—'}.")
    for item in _items(result.get("limitations"), "Существенные ограничения документа не зафиксированы."):
        _add_bullet(document, item)
    disclaimer = document.add_paragraph()
    run = disclaimer.add_run("Отчет сформирован FinSight AI. Он не является аудиторским заключением или инвестиционной рекомендацией.")
    run.italic = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(102, 119, 119)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _register_pdf_fonts() -> tuple[str, str]:
    candidates = (
        ("/System/Library/Fonts/Supplemental/Arial.ttf", "/System/Library/Fonts/Supplemental/Arial Bold.ttf"),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ("/Library/Fonts/Arial.ttf", "/Library/Fonts/Arial Bold.ttf"),
    )
    for regular, bold in candidates:
        if Path(regular).exists() and Path(bold).exists():
            try:
                pdfmetrics.registerFont(TTFont("FinSight", regular))
                pdfmetrics.registerFont(TTFont("FinSight-Bold", bold))
                return "FinSight", "FinSight-Bold"
            except Exception:
                continue
    return "Helvetica", "Helvetica-Bold"


def _p(value: Any, style: ParagraphStyle) -> Paragraph:
    text = html.escape(str(value or "—")).replace("\n", "<br/>")
    return Paragraph(text, style)


def build_pdf(result: dict[str, Any]) -> bytes:
    regular_font, bold_font = _register_pdf_fonts()
    buffer = io.BytesIO()
    page_width, _ = landscape(A4)
    document = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        rightMargin=14 * mm,
        leftMargin=14 * mm,
        topMargin=20 * mm,
        bottomMargin=16 * mm,
        title=f"FinSight AI — {result.get('metadata', {}).get('company') or 'финансовый анализ'}",
        author="FinSight AI",
    )
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle("FinTitle", parent=base["Title"], fontName=bold_font, fontSize=24, leading=29, textColor=colors.HexColor(f"#{BRAND_DARK}"), alignment=TA_LEFT, spaceAfter=7 * mm),
        "h1": ParagraphStyle("FinH1", parent=base["Heading1"], fontName=bold_font, fontSize=16, leading=20, textColor=colors.HexColor(f"#{BRAND_DARK}"), spaceBefore=5 * mm, spaceAfter=3 * mm),
        "h2": ParagraphStyle("FinH2", parent=base["Heading2"], fontName=bold_font, fontSize=12, leading=15, textColor=colors.HexColor(f"#{BRAND_TEAL}"), spaceBefore=4 * mm, spaceAfter=2 * mm),
        "body": ParagraphStyle("FinBody", parent=base["BodyText"], fontName=regular_font, fontSize=9.3, leading=13, textColor=colors.HexColor(f"#{TEXT_DARK}"), spaceAfter=2 * mm),
        "bullet": ParagraphStyle("FinBullet", parent=base["BodyText"], fontName=regular_font, fontSize=8.8, leading=12.5, leftIndent=5 * mm, firstLineIndent=-3 * mm, bulletIndent=1.5 * mm, textColor=colors.HexColor(f"#{TEXT_DARK}"), spaceAfter=1.5 * mm),
        "small": ParagraphStyle("FinSmall", parent=base["BodyText"], fontName=regular_font, fontSize=7.2, leading=9, textColor=colors.HexColor("#526767")),
        "table": ParagraphStyle("FinTable", parent=base["BodyText"], fontName=regular_font, fontSize=7.1, leading=8.7, textColor=colors.HexColor(f"#{TEXT_DARK}")),
        "table_head": ParagraphStyle("FinTableHead", parent=base["BodyText"], fontName=bold_font, fontSize=7.1, leading=8.7, textColor=colors.white, alignment=TA_CENTER),
    }

    metadata, analysis = result.get("metadata", {}), result.get("analysis", {})
    story: list[Any] = []
    story.append(_p("ПРОФЕССИОНАЛЬНЫЙ ФИНАНСОВЫЙ АНАЛИЗ", styles["h2"]))
    story.append(_p(metadata.get("company") or "Финансовый отчет", styles["title"]))
    fact_rows = [
        [_p("Исходный файл", styles["table_head"]), _p(metadata.get("filename"), styles["table"])],
        [_p("Стандарт / периметр", styles["table_head"]), _p(" · ".join(filter(None, [str(metadata.get("accounting_standard") or ""), str(metadata.get("reporting_scope") or "")])), styles["table"])],
        [_p("Отчетный год", styles["table_head"]), _p(metadata.get("reporting_year"), styles["table"])],
        [_p("Статус валидации", styles["table_head"]), _p(result.get("validation", {}).get("status"), styles["table"])],
    ]
    facts = Table(fact_rows, colWidths=[48 * mm, 140 * mm], hAlign="LEFT")
    facts.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{BRAND_DARK}")),
        ("BACKGROUND", (1, 0), (1, -1), colors.HexColor("#F2F7F7")),
        ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#BFD0D0")),
        ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D6E2E2")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.extend([facts, Spacer(1, 4 * mm), _p("Итоговое резюме", styles["h1"]), _p(analysis.get("executive_summary") or "Недостаточно подтвержденных данных для итогового резюме.", styles["body"])])
    for title, key in _analysis_sections():
        content = [_p(title, styles["h2"])]
        for item in _items(analysis.get(key)):
            content.append(Paragraph(f"•&nbsp;&nbsp;{html.escape(item)}", styles["bullet"]))
        story.append(KeepTogether(content))

    story.extend([PageBreak(), _p("Финансовая модель", styles["h1"])])
    years = _years(result)
    model_rows = [[_p(value, styles["table_head"]) for value in ["Код", "Показатель", *years, "Единица", "Источник"]]]
    for item in result.get("financial_metrics", {}).values():
        model_rows.append([
            _p(item.get("row_code") or "", styles["table"]),
            _p(item.get("name"), styles["table"]),
            *[_p(_fmt(item.get("values", {}).get(year)), styles["table"]) for year in years],
            _p(item.get("unit") or "", styles["table"]),
            _p(", ".join(map(str, item.get("source_pages", []))) or "—", styles["table"]),
        ])
    available = page_width - 28 * mm
    fixed = [18 * mm, 66 * mm, *([31 * mm] * len(years)), 25 * mm, 20 * mm]
    scale = available / sum(fixed)
    model = Table(model_rows, colWidths=[value * scale for value in fixed], repeatRows=1, hAlign="LEFT")
    model.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BRAND_DARK}")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F8F8")]),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#C9D8D8")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ALIGN", (2, 1), (1 + len(years), -1), "RIGHT"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(model)

    story.extend([PageBreak(), _p("Финансовые коэффициенты", styles["h1"])])
    ratio_rows = [[_p(value, styles["table_head"]) for value in ["Коэффициент", "Значение", "Статус", "Формула", "Интерпретация"]]]
    for item in result.get("ratios", []):
        ratio_rows.append([_p(item.get("name"), styles["table"]), _p(item.get("display"), styles["table"]), _p(item.get("status"), styles["table"]), _p(item.get("formula"), styles["table"]), _p(item.get("explanation"), styles["table"])])
    ratios = Table(ratio_rows, colWidths=[42 * mm, 25 * mm, 24 * mm, 72 * mm, available - 163 * mm], repeatRows=1)
    ratios.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{BRAND_DARK}")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F3F8F8")]),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#C9D8D8")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(ratios)
    story.extend([_p("Валидация и ограничения", styles["h1"]), _p(f"Статус автоматической валидации: {result.get('validation', {}).get('status') or '—'}.", styles["body"])])
    for item in _items(result.get("limitations"), "Существенные ограничения документа не зафиксированы."):
        story.append(Paragraph(f"•&nbsp;&nbsp;{html.escape(item)}", styles["bullet"]))
    story.append(_p("Отчет сформирован FinSight AI. Он не является аудиторским заключением или инвестиционной рекомендацией.", styles["small"]))

    def decorate(canvas, doc) -> None:
        canvas.saveState()
        width, height = landscape(A4)
        canvas.setFillColor(colors.HexColor(f"#{BRAND_DARK}"))
        canvas.rect(0, height - 12 * mm, width, 12 * mm, fill=1, stroke=0)
        canvas.setFont(bold_font, 8)
        canvas.setFillColor(colors.white)
        canvas.drawString(14 * mm, height - 8 * mm, "FINSIGHT AI  |  ФИНАНСОВЫЙ УНИВЕРСИТЕТ")
        canvas.setFont(regular_font, 7)
        canvas.setFillColor(colors.HexColor("#617777"))
        canvas.drawString(14 * mm, 8 * mm, "Автоматический финансовый анализ · результаты требуют профессионального суждения")
        canvas.drawRightString(width - 14 * mm, 8 * mm, f"Страница {doc.page}")
        canvas.restoreState()

    document.build(story, onFirstPage=decorate, onLaterPages=decorate)
    return buffer.getvalue()
