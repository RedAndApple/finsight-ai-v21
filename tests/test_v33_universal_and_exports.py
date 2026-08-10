import io
import fitz
from openpyxl import load_workbook
from docx import Document

from app.analysis import _select_financial_model
from app.exports import build_docx, build_pdf, build_xlsx
from app.financial import calculate_ratios


def _candidate(key, values, source="ras_coordinate_ocr", code=None):
    return {
        "key": key, "name": key, "unit": "тыс. руб.", "values": values,
        "source_pages": [8], "source_type": source, "confidence": 0.98,
        "row_code": code,
    }


def _ras_branch():
    values = {
        "noncurrent_assets": (400, 380), "current_assets": (600, 520), "assets": (1000, 900),
        "equity": (500, 450), "longterm_liabilities": (200, 180), "current_liabilities": (300, 270),
        "revenue": (1200, 1000), "gross_profit": (360, 300), "net_profit": (120, 90),
        "cash": (100, 80), "inventory": (150, 130), "receivables": (250, 220),
        "operating_cash_flow": (180, 150),
    }
    codes = {"assets": "1600", "equity": "1300", "current_assets": "1200", "current_liabilities": "1500", "revenue": "2110", "net_profit": "2400"}
    return [_candidate(key, {"2025": cur, "2024": prev}, code=codes.get(key)) for key, (cur, prev) in values.items()]


def _result():
    metrics, model, validation, selected, _diagnostics = _select_financial_model({
        "candidate_branches": {"ras": _ras_branch(), "ifrs": []}, "financial_candidates": [],
    })
    return {
        "id": "test", "metadata": {"company": "ПАО «Тест»", "filename": "test.pdf", "accounting_standard": "РСБУ", "reporting_year": 2025, "reporting_scope": "Отдельное юридическое лицо"},
        "financial_metrics": metrics, "canonical_financial_model": model, "validation": validation,
        "ratios": calculate_ratios(metrics), "limitations": [],
        "analysis": {"executive_summary": "Компания сохраняет устойчивую финансовую модель.", "strengths": ["Ликвидность подтверждена."], "weaknesses": [], "risks": [], "management_actions": ["Контролировать оборотный капитал."]},
    }, selected


def test_parallel_standard_selection_uses_validated_primary_model():
    result, selected = _result()
    assert selected == "ras"
    assert result["validation"]["status"] == "passed"
    assert len(result["financial_metrics"]) >= 12
    assert sum(item["status"] != "na" for item in result["ratios"]) >= 10


def test_three_professional_exports_are_openable():
    result, _ = _result()
    xlsx = build_xlsx(result)
    workbook = load_workbook(io.BytesIO(xlsx), read_only=True)
    assert {"Резюме", "Финансовая модель", "Коэффициенты", "Валидация", "Источники", "Методология"} <= set(workbook.sheetnames)

    docx = build_docx(result)
    document = Document(io.BytesIO(docx))
    content = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "ПАО «Тест»" in content
    assert "Финансовые коэффициенты" in content
    assert len(document.tables) >= 3

    pdf = build_pdf(result)
    assert pdf.startswith(b"%PDF")
    document = fitz.open(stream=pdf, filetype="pdf")
    assert len(document) >= 1
    extracted = "".join(page.get_text() for page in document).replace("\u00a0", " ")
    assert "FinSight AI" in extracted
